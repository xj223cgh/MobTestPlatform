"""AI 异步任务：触发生成用例、查询状态与结果。"""
from flask import Blueprint, request, current_app
from flask_login import login_required, current_user
from app.models.models import db, TestSuite, TestCase, Project
from app.utils.helpers import success_response, error_response
from app.utils.task_manager import task_manager, TaskStatus
import requests
import json
import re
import os

try:
    import json_repair
except ImportError:
    json_repair = None
import base64
import logging
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)

bp = Blueprint('ai_tasks', __name__, url_prefix='/api/ai-tasks')


# ---------------------------------------------------------------------------
# AI 配置加载
# ---------------------------------------------------------------------------

_ai_config_cache = {}
_ai_config_mtime = 0.0

def load_ai_config() -> dict:
    """加载 ai_config.yaml 配置，基于文件修改时间做缓存。"""
    global _ai_config_cache, _ai_config_mtime
    config_path = Path(__file__).resolve().parent.parent / 'ai' / 'ai_config.yaml'
    if not config_path.exists():
        return {}
    mtime = config_path.stat().st_mtime
    if _ai_config_cache and _ai_config_mtime == mtime:
        return _ai_config_cache
    import yaml
    with open(config_path, 'r', encoding='utf-8') as f:
        _ai_config_cache = yaml.safe_load(f) or {}
    _ai_config_mtime = mtime
    return _ai_config_cache


# ---------------------------------------------------------------------------
# 文档 & Excel 存储辅助函数
# ---------------------------------------------------------------------------

_AI_WORKSPACE = Path(__file__).resolve().parent.parent / 'ai' / 'workspace'


def _sanitize_filename(name: str, max_len: int = 50) -> str:
    """移除文件名中的不安全字符，截断到 max_len 长度。"""
    unsafe = r'<>:"/\|?*'
    clean = ''.join(c for c in name if c not in unsafe).strip()
    return clean[:max_len] or 'unnamed'


def _save_requirement_document(document_content: str, folder_name: str):
    """将原始需求文档保存到 ai/workspace/requirements/original/ 目录。"""
    try:
        req_dir = _AI_WORKSPACE / 'requirements' / 'original'
        req_dir.mkdir(parents=True, exist_ok=True)
        filepath = req_dir / f"{folder_name}.txt"
        filepath.write_text(document_content, encoding='utf-8')
        logger.info("Saved requirement document: %s", filepath)
        return str(filepath)
    except Exception as e:
        logger.warning("Failed to save requirement document: %s", e)
        return None


def _build_requirement_folder_name(suite_name: str, project_name: str = '') -> str:
    """生成需求文件夹名前缀（项目_用例集_时间戳），确保每次生成会话唯一。"""
    project_part = _sanitize_filename(project_name, 20) if project_name else 'project'
    suite_part = _sanitize_filename(suite_name, 30)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"{project_part}_{suite_part}_{ts}"


def _save_requirement_images(
    uploaded_images: list,
    docx_content: bytes | None,
    folder_name: str,
) -> str | None:
    """将所有图片（用户上传 + docx 内嵌）保存到 ai/workspace/requirements/images/{folder_name}/ 目录。"""
    all_images: list[tuple[str, bytes]] = []

    for img in (uploaded_images or []):
        fn = img.get('filename', 'image.png')
        content = img.get('content', b'')
        if content:
            all_images.append((fn, content))

    if docx_content:
        all_images.extend(_extract_images_from_docx(docx_content))

    if not all_images:
        return None

    try:
        img_dir = _AI_WORKSPACE / 'requirements' / 'images' / folder_name
        img_dir.mkdir(parents=True, exist_ok=True)

        seen_names: dict[str, int] = {}
        for name, data in all_images:
            if name in seen_names:
                seen_names[name] += 1
                stem, ext = os.path.splitext(name)
                name = f"{stem}_{seen_names[name]}{ext}"
            else:
                seen_names[name] = 0
            (img_dir / name).write_bytes(data)

        logger.info("Saved %d images to %s", len(all_images), img_dir)
        return str(img_dir)
    except Exception as e:
        logger.warning("Failed to save requirement images: %s", e)
        return None


def _save_generated_cases_excel(saved_cases: list, suite_name: str, project_name: str = ''):
    """将 AI 生成的用例导出为 Excel 文件，保存到 workspace/outputs/excel/ 目录。"""
    try:
        from app.ai.excel_exporter import export_cases_to_excel
        output_dir = _AI_WORKSPACE / 'outputs' / 'excel'
        label = f"{_sanitize_filename(project_name, 20)}_{_sanitize_filename(suite_name, 30)}" \
                if project_name else _sanitize_filename(suite_name, 30)
        result = export_cases_to_excel(saved_cases, suite_name=label, output_dir=output_dir)
        return str(result) if result else None
    except Exception as e:
        logger.warning("Failed to save cases Excel: %s", e)
        return None


# ---------------------------------------------------------------------------
# 核心任务：生成测试用例（支持单次生成和 Map-Reduce）
# ---------------------------------------------------------------------------

def generate_test_cases_task(suite_id: int, params: dict, task_manager, task_id: str):
    """在后台线程中调用 AI 生成测试用例并持久化。支持单次生成和 Map-Reduce 长文档分段生成。"""
    app = params.pop('_app', None)
    if not app:
        raise RuntimeError("缺少应用上下文，无法在后台执行任务")

    with app.app_context():
        try:
            _ensure_env_loaded()
            task_manager.update_task_status(task_id, message='正在解析需求文档...', progress=3)

            document_content = (params.get('documentContent') or '').strip()
            ai_config = get_ai_config()

            if params.get('_docx_content'):
                docx_plain = _extract_plain_text_from_docx(params['_docx_content']).strip()
                if docx_plain:
                    document_content = (
                        (document_content + '\n\n' + docx_plain).strip()
                        if document_content
                        else docx_plain
                    )

            image_text = _process_uploaded_images(params.get('_image_files', []), ai_config)
            docx_image_text = ''
            if params.get('_docx_content'):
                task_manager.update_task_status(task_id, message='正在识别文档内嵌图片...', progress=5)
                docx_image_text = _extract_and_recognize_docx_images(params['_docx_content'], ai_config)

            extra_image_text = '\n\n'.join(filter(None, [image_text, docx_image_text]))
            if extra_image_text:
                document_content = document_content + '\n\n【以下为图片识别内容】\n' + extra_image_text

            task_manager.update_task_status(task_id, message='正在检索知识库...', progress=8)
            knowledge_context = _retrieve_knowledge_context(document_content)

            test_cases = _generate_cases_from_document(
                document_content, knowledge_context, ai_config, task_manager, task_id
            )

            if not test_cases:
                raise Exception("AI未生成任何测试用例")

            task_manager.update_task_status(
                task_id,
                message=f'正在保存测试用例，共{len(test_cases)}条...',
                progress=60
            )

            suite = TestSuite.query.get(suite_id)
            if not suite:
                raise Exception("用例集不存在")
            project_id = params.get('projectId') or suite.project_id
            iteration_id = params.get('iterationId') or suite.iteration_id
            version_requirement_id = params.get('requirementId') or suite.version_requirement_id
            if project_id is None:
                first_project = Project.query.first()
                if first_project:
                    project_id = first_project.id
            if project_id is None:
                raise ValueError("用例集未关联项目，且系统中无可用项目。请先在「项目」中创建项目，并为用例集选择所属项目。")
            creator_id = params.get('creatorId') or getattr(suite, 'creator_id', None)
            if creator_id is None:
                raise ValueError("无法确定用例创建人，请重新发起生成任务。")

            case_number_prefix = generate_case_number_prefix(suite, params)
            max_index = get_max_case_index(suite_id)

            saved_cases = []
            total_cases = len(test_cases)

            for i, case_item in enumerate(test_cases):
                suite = TestSuite.query.get(suite_id)
                if not suite:
                    task_manager.update_task_status(
                        task_id,
                        status=TaskStatus.FAILED,
                        message='用例集已删除，生成已终止',
                        completed_at=datetime.now().isoformat()
                    )
                    break
                current_progress = 60 + int((i / total_cases) * 35)
                task_manager.update_task_status(
                    task_id,
                    message=f'正在保存第{i+1}/{total_cases}条用例...',
                    progress=current_progress,
                    current=i + 1,
                    total=total_cases
                )

                current_index = max_index + i + 1
                suffix = str(current_index).zfill(3)
                if current_index > 999:
                    suffix = "999"
                case_number = f"{case_number_prefix}{suffix}"

                case_data = {
                    'suite_id': suite_id,
                    'case_number': case_number,
                    'case_name': case_item.get('case_name', f'测试用例_{case_number}'),
                    'case_description': case_item.get('case_description', ''),
                    'priority': case_item.get('priority', 'P1'),
                    'status': case_item.get('status', ''),
                    'preconditions': case_item.get('preconditions', ''),
                    'steps': case_item.get('steps', ''),
                    'expected_result': case_item.get('expected_result', ''),
                    'test_data': case_item.get('test_data', ''),
                    'project_id': project_id,
                    'iteration_id': iteration_id,
                    'version_requirement_id': version_requirement_id,
                    'creator_id': creator_id,
                }

                test_case = TestCase(**case_data)
                db.session.add(test_case)
                saved_cases.append(case_data)

            if suite:
                suite.case_count = len(saved_cases)
            db.session.commit()

            project_obj = Project.query.get(project_id) if project_id else None
            project_label = project_obj.project_name if project_obj else ''
            suite_label_for_file = suite.suite_name if suite else 'suite'
            folder_name = _build_requirement_folder_name(suite_label_for_file, project_label)
            _save_requirement_document(document_content, folder_name)
            _save_requirement_images(params.get('_image_files', []), params.get('_docx_content'), folder_name)
            _save_generated_cases_excel(saved_cases, suite_label_for_file, project_label)

            debug_mode = params.get('debugMode', False)
            if not debug_mode:
                try:
                    _ensure_env_loaded()
                    if os.getenv('EMBEDDING_MODEL'):
                        from app.services.knowledge_service import upload_document
                        suite_label = suite.suite_name if suite else '需求文档'
                        kb_filename = f"{suite_label}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                        upload_document(
                            file_content=document_content.encode('utf-8'),
                            filename=kb_filename,
                            metadata={'source': 'auto_generate', 'suite_name': suite_label},
                        )
                        logger.info("Auto-stored requirement to knowledge base: %s", kb_filename)
                except Exception as kb_err:
                    logger.warning("Auto-store to knowledge base failed (non-fatal): %s", kb_err)

            msg_suffix = '（调试模式，未存入知识库）' if debug_mode else ''
            task_manager.update_task_status(
                task_id,
                message=f'成功生成并保存{len(saved_cases)}条测试用例{msg_suffix}',
                progress=100
            )
            notify_creator_id = params.get('creatorId') or creator_id
            if notify_creator_id:
                from app.services.notification_service import notify_users
                suite_label = suite.suite_name if suite else '用例集'
                notify_users(
                    [notify_creator_id], 'ai_case_generated', 'AI 用例生成完成',
                    f'用例集「{suite_label}」的 AI 用例生成已完成，共生成并保存 {len(saved_cases)} 条测试用例',
                    'suite', suite_id, extra={'total_cases': len(saved_cases)}
                )
            return {
                'suite_id': suite_id,
                'total_cases': len(saved_cases),
                'saved_cases': saved_cases[:5]
            }

        except Exception as e:
            db.session.rollback()
            notify_creator_id = params.get('creatorId')
            if notify_creator_id is None:
                s = TestSuite.query.get(suite_id)
                if s:
                    notify_creator_id = getattr(s, 'creator_id', None)
            if notify_creator_id:
                try:
                    from app.services.notification_service import notify_users
                    suite_obj = TestSuite.query.get(suite_id)
                    suite_label = suite_obj.suite_name if suite_obj else '用例集'
                    notify_users(
                        [notify_creator_id], 'ai_case_generated', 'AI 用例生成失败',
                        f'用例集「{suite_label}」的 AI 用例生成失败：{str(e)[:150]}',
                        'suite', suite_id, extra={'error': str(e)[:200]}
                    )
                except Exception:
                    pass
            raise


# ---------------------------------------------------------------------------
# 知识检索（Phase 3 — RAG）
# ---------------------------------------------------------------------------

def _retrieve_knowledge_context(document_content: str) -> str:
    """从知识库检索与需求文档相关的背景知识，参数取自 ai_config.yaml 中的 retrieval 配置。"""
    try:
        from app.services.knowledge_service import search
        if not document_content or not document_content.strip():
            return ''

        cfg = load_ai_config().get('knowledge', {}).get('retrieval', {})
        query_max = cfg.get('query_max_length', 500)
        top_k = cfg.get('top_k', 5)
        threshold = cfg.get('similarity_threshold', 0.5)

        query = document_content[:query_max]
        results = search(query, top_k=top_k)
        if not results:
            return ''
        relevant = [r for r in results if r.get('distance', 1.0) < threshold]
        if not relevant:
            return ''
        parts = []
        for r in relevant[:top_k]:
            filename = r.get('metadata', {}).get('filename', '')
            category = r.get('metadata', {}).get('category_label', '')
            text = r.get('text', '')
            source_label = f"{category} - {filename}" if category else filename
            if source_label:
                parts.append(f"[来源: {source_label}]\n{text}")
            else:
                parts.append(text)
        return '\n\n---\n\n'.join(parts)
    except Exception:
        return ''


# ---------------------------------------------------------------------------
# 图片提取 & 识别（Vision API）
# ---------------------------------------------------------------------------

_SUPPORTED_IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}


def _extract_plain_text_from_docx(file_content: bytes) -> str:
    """从 .docx 中提取正文段落为纯文本（与知识库 .docx 解析一致）。"""
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_content))
        return '\n\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.warning('Failed to extract plain text from docx: %s', e)
        return ''


def _extract_images_from_docx(file_content: bytes) -> list:
    """从 .docx 文件中提取内嵌图片，返回 (filename, bytes) 元组列表。"""
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_content))
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_part = rel.target_part
                img_bytes = img_part.blob
                img_name = img_part.partname.split('/')[-1]
                images.append((img_name, img_bytes))
        return images
    except Exception as e:
        logger.warning("Failed to extract images from docx: %s", e)
        return []


def _recognize_image_via_vision_api(image_bytes: bytes, image_filename: str, ai_config: dict) -> str:
    """调用视觉大模型识别图片内容（用于提取需求图片中的文字与关键信息）。"""
    _ensure_env_loaded()
    vision_model = (os.getenv('AI_VISION_MODEL') or '').strip()
    if not vision_model:
        return ''

    api_key = (ai_config.get('apiKey') or '').strip()
    base_url = (ai_config.get('baseURL') or '').strip().rstrip('/')
    if not api_key or api_key.startswith('sk-your'):
        return ''

    ext = os.path.splitext(image_filename)[1].lower()
    mime_map = {'.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                '.gif': 'image/gif', '.bmp': 'image/bmp', '.webp': 'image/webp'}
    mime_type = mime_map.get(ext, 'image/png')

    b64 = base64.b64encode(image_bytes).decode('ascii')
    data_url = f"data:{mime_type};base64,{b64}"

    url = f"{base_url}/chat/completions"
    headers = {'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'}
    payload = {
        'model': vision_model,
        'messages': [
            {'role': 'user', 'content': [
                {'type': 'text', 'text': '请详细描述这张图片中的所有文字内容和关键信息，用于软件测试需求分析。如果是流程图或界面截图，请描述其结构和交互元素。'},
                {'type': 'image_url', 'image_url': {'url': data_url}}
            ]}
        ],
        'max_tokens': 1024,
    }
    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=60)
        resp.raise_for_status()
        result = resp.json()
        content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
        return content.strip()
    except Exception as e:
        logger.warning("Vision API call failed for %s: %s", image_filename, e)
        return ''


def _process_uploaded_images(image_files: list, ai_config: dict) -> str:
    """将用户上传的图片逐张调用视觉 API 识别，返回拼合后的文本。"""
    if not image_files:
        return ''
    parts = []
    for img_file in image_files:
        filename = img_file.get('filename', 'image.png')
        content = img_file.get('content', b'')
        if not content:
            continue
        text = _recognize_image_via_vision_api(content, filename, ai_config)
        if text:
            parts.append(f"[图片: {filename}]\n{text}")
    return '\n\n'.join(parts)


def _extract_and_recognize_docx_images(docx_content: bytes, ai_config: dict) -> str:
    """从 .docx 中提取图片并调用视觉 API 识别，返回拼合后的文本。"""
    images = _extract_images_from_docx(docx_content)
    if not images:
        return ''
    try:
        max_images = int(os.getenv('AI_VISION_MAX_IMAGES', '0'))
    except (TypeError, ValueError):
        max_images = 0
    if max_images <= 0:
        max_images = len(images)
    parts = []
    for img_name, img_bytes in images[:max_images]:
        text = _recognize_image_via_vision_api(img_bytes, img_name, ai_config)
        if text:
            parts.append(f"[文档内嵌图片: {img_name}]\n{text}")
    return '\n\n'.join(parts)



# ---------------------------------------------------------------------------
# 生成：单次生成 vs Map-Reduce
# ---------------------------------------------------------------------------

def _generate_cases_from_document(document_content, knowledge_context, ai_config, task_manager, task_id):
    """生成测试用例；文档超过阈值时自动切换为 Map-Reduce 分段生成。"""
    doc_cfg = load_ai_config().get('document_processing', {})
    map_reduce_threshold = doc_cfg.get('map_reduce_threshold', 3000)
    doc_len = len(document_content or '')

    if doc_len >= map_reduce_threshold:
        return _generate_map_reduce(
            document_content, knowledge_context, ai_config, task_manager, task_id
        )

    task_manager.update_task_status(task_id, message='正在构建AI提示词...', progress=20)
    prompt = build_test_case_prompt({}, document_content, knowledge_context)

    task_manager.update_task_status(task_id, message='正在调用AI生成用例...', progress=30)
    base_max = ai_config.get('maxTokens', 4096)
    extra_tokens = min(12288, (doc_len // 1000) * 500)
    dynamic_max_tokens = base_max + extra_tokens

    cases = _call_and_parse_with_retry(prompt, ai_config, dynamic_max_tokens)

    task_manager.update_task_status(task_id, message='正在校验用例字段...', progress=55)
    return [_validate_and_fix_case(c) for c in cases]


def _generate_map_reduce(document_content, knowledge_context, ai_config, task_manager, task_id):
    """Map-Reduce 模式：拆分文档 → 逐段生成用例 → 合并去重。"""
    from app.utils.doc_chunker import split_document, extract_summary
    from app.utils.prompt_loader import render_prompt

    task_manager.update_task_status(task_id, message='正在拆分需求文档...', progress=10)
    chunks = split_document(document_content)
    summary = extract_summary(document_content)
    total_suggested = _suggest_case_count(document_content)

    if not chunks:
        raise Exception("文档拆分失败，无可处理段落")

    all_cases = []
    num_chunks = len(chunks)
    per_chunk_min = max(3, total_suggested // num_chunks)
    per_chunk_max = per_chunk_min + 5

    for i, chunk in enumerate(chunks):
        progress = 12 + int((i / num_chunks) * 70)
        task_manager.update_task_status(
            task_id,
            message=f'正在生成第 {i + 1}/{num_chunks} 段用例...',
            progress=progress
        )

        try:
            prompt = render_prompt(
                'generate_cases_chunk.yaml',
                chunk_content=chunk,
                chunk_index=i + 1,
                total_chunks=num_chunks,
                document_summary=summary,
                knowledge_context=knowledge_context,
                suggested_min=per_chunk_min,
                suggested_max=per_chunk_max,
            )
        except Exception:
            prompt = _build_chunk_fallback_prompt(
                chunk, i + 1, num_chunks, summary,
                knowledge_context, per_chunk_min, per_chunk_max,
            )

        chunk_len = len(chunk)
        base_max = ai_config.get('maxTokens', 4096)
        extra = min(12288, (chunk_len // 1000) * 500)
        dynamic_max = base_max + extra

        try:
            cases = _call_and_parse_with_retry(prompt, ai_config, dynamic_max)
            all_cases.extend(_validate_and_fix_case(c) for c in cases)
        except Exception as e:
            task_manager.update_task_status(
                task_id,
                message=f'第 {i + 1}/{num_chunks} 段生成失败({str(e)[:80]})，继续下一段...',
            )

    task_manager.update_task_status(task_id, message='正在合并去重用例...', progress=85)
    deduplicated = _deduplicate_cases(all_cases)

    task_manager.update_task_status(
        task_id,
        message=f'分段生成完成，共 {len(deduplicated)} 条用例（去重前 {len(all_cases)} 条）',
        progress=58
    )
    return deduplicated


def _deduplicate_cases(cases: list, threshold: float = 0.85) -> list:
    """按用例名称相似度去重，阈值 ≥ threshold 视为重复。"""
    from difflib import SequenceMatcher
    result = []
    for case in cases:
        name = case.get('case_name', '')
        is_dup = any(
            SequenceMatcher(None, name, existing.get('case_name', '')).ratio() >= threshold
            for existing in result
        )
        if not is_dup:
            result.append(case)
    return result


def _build_chunk_fallback_prompt(chunk, chunk_index, total_chunks, summary,
                                 knowledge_context, suggested_min, suggested_max):
    """分段模板缺失时的回退提示词构建。"""
    parts = [f"你是一名专业测试工程师。以下是需求文档的第 {chunk_index}/{total_chunks} 部分。"]
    parts.append(f"\n【文档总览】\n{summary}")
    if knowledge_context:
        parts.append(f"\n【业务背景知识（仅供参考）】\n{knowledge_context}")
    parts.append(f"\n【本段需求内容】\n{chunk}")
    parts.append(f"""
【输出要求】
1. 以 JSON 格式返回，且只返回 JSON。
2. 格式如下：
{{
  "test_cases": [
    {{
      "case_name": "用例名称",
      "case_description": "用例描述",
      "priority": "P0/P1/P2/P3/P4",
      "preconditions": "前置条件",
      "steps": "测试步骤",
      "expected_result": "预期结果",
      "test_data": "测试数据"
    }}
  ]
}}

【质量与数量要求】
- 建议本段生成约 {suggested_min}～{suggested_max} 条用例。
- 只输出上述 JSON。""")
    return '\n'.join(parts)


# ---------------------------------------------------------------------------
# 重试 + 字段校验
# ---------------------------------------------------------------------------

_MAX_RETRIES = 1


def _call_and_parse_with_retry(prompt: str, ai_config: dict, max_tokens: int) -> list:
    """调用 AI API 并解析返回的 JSON；若解析失败则重试一次。"""
    last_error = None
    for attempt in range(_MAX_RETRIES + 1):
        try:
            ai_response = call_ai_api(prompt, ai_config, max_tokens_override=max_tokens)
            return parse_ai_response(ai_response)
        except Exception as e:
            err_msg = str(e)
            is_parse_error = '解析' in err_msg or 'JSON' in err_msg
            if is_parse_error and attempt < _MAX_RETRIES:
                last_error = e
                continue
            raise
    raise last_error


_VALID_PRIORITIES = {'P0', 'P1', 'P2', 'P3', 'P4'}

_PRIORITY_MAP = {
    '高': 'P0', '最高': 'P0', 'high': 'P0', 'critical': 'P0',
    '中': 'P1', 'medium': 'P1', 'normal': 'P1',
    '低': 'P2', 'low': 'P2',
    '最低': 'P3', 'minor': 'P3',
    'trivial': 'P4',
}


def _validate_and_fix_case(case: dict) -> dict:
    """校验并修正单条用例：规范化优先级、填充缺失字段、将列表值转为换行文本。"""
    priority = (case.get('priority') or 'P1').strip()
    if priority not in _VALID_PRIORITIES:
        priority = _PRIORITY_MAP.get(priority.lower(), _PRIORITY_MAP.get(priority, 'P1'))
    case['priority'] = priority

    case.setdefault('case_name', '未命名用例')
    case.setdefault('case_description', '')
    case.setdefault('preconditions', '')
    case.setdefault('steps', '')
    case.setdefault('expected_result', '')
    case.setdefault('test_data', '')

    for field in ('case_name', 'case_description', 'preconditions', 'steps', 'expected_result', 'test_data'):
        val = case.get(field)
        if isinstance(val, list):
            case[field] = '\n'.join(str(item) for item in val)
        elif val is not None:
            case[field] = str(val)

    return case


# ---------------------------------------------------------------------------
# 提示词构建
# ---------------------------------------------------------------------------

def _suggest_case_count(document_content: str) -> int:
    """根据文档长度估算建议用例条数，参数从 ai_config.yaml 读取。"""
    cfg = load_ai_config().get('document_processing', {}).get('case_count_estimation', {})
    min_per_1000 = cfg.get('min_per_1000_chars', 3)
    base_min = cfg.get('base_minimum', 8)
    max_cap = cfg.get('max_cap', 80)

    text_len = len((document_content or '').strip())
    if text_len <= 0:
        return base_min
    return max(base_min, min(max_cap, (text_len // 1000) * min_per_1000 * 3))


def build_test_case_prompt(params: dict, document_content: str, knowledge_context: str = '') -> str:
    """构建 AI 提示词（从模板加载，失败时回退硬编码）。"""
    suggested_count = _suggest_case_count(document_content)
    doc_content = (document_content or '').strip() or "（未提供需求文档内容）"

    try:
        from app.utils.prompt_loader import render_prompt
        return render_prompt(
            'generate_cases.yaml',
            document_content=doc_content,
            knowledge_context=knowledge_context,
            suggested_min=suggested_count,
            suggested_max=suggested_count + 10,
        )
    except Exception:
        return _build_fallback_prompt(doc_content, suggested_count, knowledge_context)


def _build_fallback_prompt(doc_content: str, suggested_count: int, knowledge_context: str = '') -> str:
    """模板不可用时的硬编码回退提示词。"""
    knowledge_block = ''
    if knowledge_context:
        knowledge_block = f"""
【业务背景知识（仅供参考，不作为用例依据）】
{knowledge_context}
"""
    return f"""你是一名专业测试工程师。请**严格依据**下方「需求文档内容」生成功能测试用例。
每条用例的步骤、预期结果必须能在需求文档中找到对应依据，不要编造需求中未提及的行为。
{knowledge_block}
【需求文档内容】
{doc_content}

【输出要求】
1. 以 JSON 格式返回，且只返回 JSON，不要有任何其他说明文字。
2. 格式如下：
{{
  "test_cases": [
    {{
      "case_name": "用例名称",
      "case_description": "用例描述",
      "priority": "P0/P1/P2/P3/P4",
      "preconditions": "前置条件",
      "steps": "测试步骤（每步一行，可执行、可验证）",
      "expected_result": "预期结果（可验证、与需求对应）",
      "test_data": "测试数据"
    }}
  ]
}}

【质量与数量要求】
- 优先级：P0 最高，P4 最低；核心流程用 P0/P1，异常/边界用 P2/P3。
- 覆盖：对文档中的每个功能点/场景，至少生成正常流程、异常或边界类用例，不要遗漏重要功能。
- 数量：建议本需求生成约 **{suggested_count}～{suggested_count + 10}** 条用例。
- 步骤与结果：测试步骤要具体、可执行；预期结果要可验证，并与需求描述一致。
- 只输出上述 JSON，不要输出 ```json 以外的标记或解释。
"""


# ---------------------------------------------------------------------------
# AI API 调用
# ---------------------------------------------------------------------------

_DEFAULT_SYSTEM_PROMPT = (
    "你是专业测试工程师。根据用户提供的需求文档生成测试用例时，必须严格依据文档内容：\n"
    "- 每条用例的步骤和预期结果需与文档中的描述对应，不编造文档未提及的功能或规则。\n"
    "- 按功能点/场景完整覆盖，需求多则用例数量应相应增加，不要人为限制在固定条数。\n"
    "- 只输出用户要求的 JSON，不要输出任何解释、代码块标记或多余文字。"
)


def call_ai_api(prompt: str, ai_config: dict, max_tokens_override: int = None,
                system_prompt: str = None) -> dict:
    """调用 Chat Completions API"""
    api_key = (ai_config.get('apiKey') or '').strip()
    if not api_key or api_key == 'sk-your-api-key' or api_key == 'sk-your-api-key-here':
        raise ValueError(
            '未配置有效的 AI_API_KEY。请在 backend/.env 中设置 AI_API_KEY，'
            '并到 SiliconFlow 控制台申请/复制密钥：https://cloud.siliconflow.cn/'
        )

    if system_prompt is None:
        try:
            from app.utils.prompt_loader import load_system_prompt
            system_prompt = load_system_prompt()
        except Exception:
            system_prompt = _DEFAULT_SYSTEM_PROMPT

    base_url = (ai_config.get('baseURL') or '').strip().rstrip('/')
    url = f"{base_url}/chat/completions"
    headers = {
        'Authorization': f"Bearer {api_key}",
        'Content-Type': 'application/json'
    }
    max_tokens = max_tokens_override if max_tokens_override is not None else ai_config.get('maxTokens', 4096)
    data = {
        'model': ai_config.get('model', 'Qwen/Qwen2.5-7B-Instruct'),
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': prompt}
        ],
        'temperature': ai_config.get('temperature', 0.3),
        'max_tokens': max_tokens
    }
    try:
        response = requests.post(url, headers=headers, json=data, timeout=120)
        if response.status_code == 401:
            raise ValueError(
                'AI 服务认证失败(401)。请检查 backend/.env 中的 AI_API_KEY 是否有效、未过期，'
                '并到 SiliconFlow 控制台确认密钥状态：https://cloud.siliconflow.cn/'
            )
        response.raise_for_status()
        resp_json = response.json()
        if isinstance(resp_json, dict) and 'error' in resp_json:
            err = resp_json['error']
            msg = err.get('message', err) if isinstance(err, dict) else str(err)
            raise ValueError(f'AI 服务返回错误: {msg}')
        return resp_json
    except requests.exceptions.HTTPError as e:
        if e.response is not None and e.response.status_code == 401:
            raise ValueError(
                'AI 服务认证失败(401)。请检查 backend/.env 中的 AI_API_KEY 是否有效、未过期。'
            ) from e
        raise


# ---------------------------------------------------------------------------
# 响应解析
# ---------------------------------------------------------------------------

def _strip_ai_markdown_json_fence(content: str) -> str:
    """去掉 ``` / ```json 包裹及首尾说明文字中的代码块。"""
    text = (content or "").strip()
    if not text:
        return text
    m = re.search(r"```(?:json)?\s*\r?\n(.*?)```", text, flags=re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return text


def _extract_first_balanced_json_object(text: str) -> str | None:
    """从首个 {{ 起截取与之配平的 JSON 对象（字符串内引号与 \\ 转义参与判断）。"""
    start = text.find("{")
    if start < 0:
        return None
    depth = 0
    in_string = False
    escape = False
    i = start
    while i < len(text):
        ch = text[i]
        if escape:
            escape = False
            i += 1
            continue
        if in_string:
            if ch == "\\":
                escape = True
            elif ch == '"':
                in_string = False
            i += 1
            continue
        if ch == '"':
            in_string = True
            i += 1
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start : i + 1]
        i += 1
    return None


def _repair_common_json_text(s: str) -> str:
    """模型常见笔误：尾随逗号、弯引号（避免打断 JSON 字符串边界）。"""
    s = s.strip()
    # 中文弯引号在字符串值内易与 ASCII " 冲突，改为单引号（JSON 字符串内允许 '）
    s = s.replace("\u201c", "'").replace("\u201d", "'")
    s = s.replace("\u2018", "'").replace("\u2019", "'")
    # }, ] 前的多余逗号
    s = re.sub(r",(\s*[\]}])", r"\1", s)
    return s


def _parse_ai_raw_json_text(raw: str):
    """严格 json.loads → 简单修复后 loads → json_repair.loads（若已安装）。"""
    variants = []
    seen = set()
    for v in (raw, _repair_common_json_text(raw)):
        if not v or v in seen:
            continue
        seen.add(v)
        variants.append(v)

    last_err: Exception | None = None
    for text in variants:
        try:
            return json.loads(text)
        except json.JSONDecodeError as e:
            last_err = e
        except Exception as e:
            last_err = e

    if json_repair is not None:
        for text in variants:
            try:
                return json_repair.loads(text)
            except Exception as e:
                last_err = e

    if isinstance(last_err, json.JSONDecodeError):
        raise last_err
    if last_err is not None:
        raise json.JSONDecodeError(str(last_err), raw, 0) from last_err
    raise json.JSONDecodeError("empty AI JSON content", raw, 0)


def parse_ai_content_to_payload(content: str) -> dict | list:
    """将模型输出的正文解析为 dict 或 list（含 test_cases 的对象，或直接为用例数组）。"""
    stripped = _strip_ai_markdown_json_fence(content)
    if not stripped:
        raise json.JSONDecodeError("empty content after strip", content or "", 0)

    candidates = [stripped]
    inner = _extract_first_balanced_json_object(stripped)
    if inner and inner != stripped:
        candidates.append(inner)

    last_err: Exception | None = None
    for cand in candidates:
        try:
            return _parse_ai_raw_json_text(cand)
        except (json.JSONDecodeError, ValueError, TypeError) as e:
            last_err = e
    if last_err:
        raise last_err
    raise json.JSONDecodeError("unable to parse AI JSON", stripped, 0)


def parse_ai_response(ai_response: dict) -> list:
    """解析 AI 返回 JSON"""
    if not isinstance(ai_response, dict):
        raise Exception("解析AI返回结果失败: 响应不是有效的 JSON 对象")
    if 'error' in ai_response:
        err = ai_response['error']
        msg = err.get('message', err) if isinstance(err, dict) else str(err)
        raise Exception(f"AI 返回错误: {msg}")
    try:
        choices = ai_response.get('choices') or []
        if not choices:
            raise Exception("AI 返回中无 choices，可能为模型限流或服务异常，请稍后重试")
        first = choices[0]
        if not isinstance(first, dict):
            raise Exception("AI 返回的 choices[0] 格式异常")
        message = first.get('message') or {}
        if not isinstance(message, dict):
            raise Exception("AI 返回的 message 格式异常")
        content = message.get('content')
        if content is None:
            raise Exception("AI 返回的 content 为空，请检查模型或提示词")
        content = str(content).strip()
        if not content:
            raise Exception("AI 返回的 content 为空字符串")

        parsed = parse_ai_content_to_payload(content)
        if isinstance(parsed, list):
            cases = parsed
        elif isinstance(parsed, dict):
            cases = parsed.get('test_cases')
            if cases is None and parsed:
                # 偶发仅返回单条用例对象
                if any(k in parsed for k in ('case_name', 'steps', 'expected_result')):
                    cases = [parsed]
                else:
                    cases = []
        else:
            cases = []
        return list(cases) if cases else []
    except json.JSONDecodeError as e:
        logger.warning("AI 返回 JSON 解析失败（首段字符预览）: %s", content[:800] if content else "")
        raise Exception(f"解析AI返回结果失败: 内容不是合法 JSON（{e}）")
    except Exception as e:
        if isinstance(e, Exception) and not isinstance(e, (KeyError, IndexError, TypeError)):
            raise
        raise Exception(f"解析AI返回结果失败: {str(e)}")


# ---------------------------------------------------------------------------
# 配置 & 环境变量辅助函数
# ---------------------------------------------------------------------------

def _ensure_env_loaded():
    """确保已从 backend/.env 加载环境变量"""
    from pathlib import Path
    _backend_dir = Path(__file__).resolve().parent.parent.parent
    _env_file = _backend_dir / '.env'
    if _env_file.exists():
        from dotenv import load_dotenv
        load_dotenv(dotenv_path=_env_file, override=True)


def get_ai_config() -> dict:
    """获取AI配置"""
    _ensure_env_loaded()
    api_key = (os.getenv('AI_API_KEY') or 'sk-your-api-key').strip()
    base_url = (os.getenv('AI_BASE_URL') or 'https://api.siliconflow.cn/v1').strip().rstrip('/')
    try:
        temperature = float(os.getenv('AI_TEMPERATURE', '0.3'))
    except (TypeError, ValueError):
        temperature = 0.3
    try:
        max_tokens = int(os.getenv('AI_MAX_TOKENS', '4096'))
    except (TypeError, ValueError):
        max_tokens = 4096
    return {
        'baseURL': base_url,
        'apiKey': api_key,
        'model': (os.getenv('AI_MODEL') or 'Qwen/Qwen2.5-7B-Instruct').strip(),
        'temperature': temperature,
        'maxTokens': max_tokens
    }


# ---------------------------------------------------------------------------
# 用例编号生成
# ---------------------------------------------------------------------------

_ZH2EN = {
    "项目": "Project", "用户": "User", "管理": "Management", "需求": "Requirement",
    "登录": "Login", "系统": "System", "测试": "Test", "模块": "Module", "功能": "Function",
    "平台": "Platform", "版本": "Version", "迭代": "Iteration", "接口": "Interface",
    "服务": "Service", "后台": "Backend", "前端": "Frontend", "数据": "Data",
    "订单": "Order", "支付": "Payment", "消息": "Message", "配置": "Config",
    "权限": "Permission", "角色": "Role", "审核": "Review", "报表": "Report",
    "文件": "File", "上传": "Upload", "下载": "Download", "搜索": "Search",
    "列表": "List", "详情": "Detail", "新增": "Create", "编辑": "Edit", "删除": "Delete",
    "移动": "Mobile", "网页": "Web", "应用": "Application", "中心": "Center",
    "个人": "Personal", "账户": "Account", "设置": "Settings", "首页": "Home",
    "单": "Single", "新": "New", "旧": "Old", "中": "Center", "心": "Core",
    "项": "Project", "目": "Item", "用": "Use", "户": "User", "需": "Requirement",
    "求": "Demand", "登": "Login", "录": "Record", "测": "Test", "试": "Test",
    "模": "Module", "块": "Block", "功": "Function", "能": "Capability",
    "系": "System", "统": "System", "版": "Version", "本": "Version",
}


def _english_word_initials(english_phrase: str) -> str:
    """提取英文短语每个单词的首字母大写拼接（如 "Test Case" → "TC"）。"""
    if not english_phrase or not english_phrase.strip():
        return ""
    return "".join(w[0].upper() for w in english_phrase.strip().split() if w and w[0].isalpha())


def _chinese_char_to_pinyin_initial(char: str) -> str:
    """将单个中文字符转为拼音首字母大写，需 pypinyin 库支持。"""
    if not char or not ("\u4e00" <= char <= "\u9fff"):
        return ""
    try:
        from pypinyin import pinyin, Style
        py = pinyin(char, style=Style.FIRST_LETTER)
        if py and py[0] and py[0][0]:
            return py[0][0].upper()
    except Exception:
        pass
    return ""


def _name_to_english_abbrev(name: str, max_len: int = 3) -> str:
    """将中英文混合名称转为英文缩写（中文先查词典再取拼音首字母），用于用例编号前缀。"""
    if not name or not str(name).strip():
        return ""
    name = str(name).strip()
    result = []
    i = 0
    while i < len(name):
        char = name[i]
        if char.isalnum() or ord(char) < 128:
            result.append(char.upper())
            i += 1
            continue
        if "\u4e00" <= char <= "\u9fff":
            two = name[i: i + 2] if i + 2 <= len(name) else ""
            one = char
            if two and two in _ZH2EN:
                initials = _english_word_initials(_ZH2EN[two])
                result.append(initials)
                i += 2
            elif one in _ZH2EN:
                initials = _english_word_initials(_ZH2EN[one])
                result.append(initials)
                i += 1
            else:
                py = _chinese_char_to_pinyin_initial(one)
                if py:
                    result.append(py)
                if two:
                    py2 = _chinese_char_to_pinyin_initial(two[1])
                    if py2:
                        result.append(py2)
                    i += 2
                else:
                    i += 1
            continue
        i += 1
    abbrev = "".join(result)[:max_len]
    return abbrev if abbrev else ""


def generate_case_number_prefix(suite, params: dict) -> str:
    """生成用例编号前缀：项目缩写-版本号-需求缩写（如 MTP-1.0.0-REQ）。"""
    import re
    project_name = params.get("projectName") or (suite.project.project_name if suite.project else "") or "PROJ"
    iteration_name = params.get("iterationName") or (suite.iteration.iteration_name if suite.iteration else "") or "1.0.0"
    requirement_name = params.get("requirementName") or (suite.version_requirement.requirement_name if suite.version_requirement else "") or "REQ"

    project_short = _name_to_english_abbrev(project_name, 3) or "PROJ"
    version_match = re.search(r"\d+\.\d+\.\d+", str(iteration_name))
    version = version_match.group(0) if version_match else "1.0.0"
    requirement_short = _name_to_english_abbrev(requirement_name, 3) or "REQ"

    return f"{project_short}-{version}-{requirement_short}"


def get_max_case_index(suite_id: int) -> int:
    """查询用例集中已有用例编号的最大序号（末尾 3 位数字），用于自动递增编号。"""
    import re
    try:
        cases = TestCase.query.filter_by(suite_id=suite_id).all()
        if not cases:
            return 0
        max_index = 0
        for case in cases:
            match = re.search(r'(\d{3})$', case.case_number or '')
            if match:
                index = int(match.group(1))
                max_index = max(max_index, index)
        return max_index
    except Exception:
        return 0


# ---------------------------------------------------------------------------
# 路由处理  
# ---------------------------------------------------------------------------

@bp.route('/generate-cases', methods=['POST'])
@login_required
def generate_cases():
    """创建 AI 生成测试用例的异步任务。支持 JSON 或 multipart/form-data（含图片）。"""
    try:
        if request.content_type and 'multipart' in request.content_type:
            data = request.form.to_dict()
            data['suite_id'] = int(data.get('suite_id', 0)) or None
        else:
            data = request.get_json()

        suite_id = data.get('suite_id')
        if not suite_id:
            return error_response(400, '缺少用例集ID')

        suite = TestSuite.query.get(suite_id)
        if not suite:
            return error_response(404, '用例集不存在')

        image_files = []
        docx_content = None
        if request.content_type and 'multipart' in request.content_type:
            for key in request.files:
                f = request.files[key]
                if not f.filename:
                    continue
                ext = os.path.splitext(f.filename)[1].lower()
                content = f.read()
                if ext in _SUPPORTED_IMAGE_EXTENSIONS:
                    image_files.append({'filename': f.filename, 'content': content})
                elif ext == '.docx':
                    docx_content = content

        params = {
            '_app': current_app._get_current_object(),
            'projectId': data.get('projectId'),
            'iterationId': data.get('iterationId'),
            'requirementId': data.get('requirementId'),
            'projectName': data.get('projectName', ''),
            'iterationName': data.get('iterationName', ''),
            'requirementName': data.get('requirementName', ''),
            'description': data.get('description', ''),
            'documentContent': data.get('documentContent', ''),
            'creatorId': current_user.id,
            'debugMode': str(data.get('debugMode', '')).lower() in ('true', '1', 'yes'),
            '_image_files': image_files,
            '_docx_content': docx_content,
        }

        task_id = task_manager.create_task(
            task_name=f'AI生成测试用例 - {suite.suite_name}',
            task_func=generate_test_cases_task,
            suite_id=suite_id,
            params=params
        )
        task_manager.update_task_status(task_id, suite_id=suite_id)

        return success_response({
            'task_id': task_id,
            'suite_id': suite_id,
            'message': '任务已创建，正在后台生成测试用例'
        })

    except Exception as e:
        return error_response(500, f'创建任务失败: {str(e)}')


@bp.route('/suite/<int:suite_id>/generating', methods=['GET'])
@login_required
def get_suite_generating(suite_id):
    """查询指定用例集是否正在AI生成中"""
    try:
        with task_manager.lock:
            for tid, task in task_manager.tasks.items():
                if task.get('suite_id') == suite_id and task.get('status') in (TaskStatus.PENDING, TaskStatus.RUNNING):
                    return success_response({'generating': True, 'task_id': tid})
        return success_response({'generating': False})
    except Exception as e:
        return error_response(500, f'查询失败: {str(e)}')


@bp.route('/task-status/<task_id>', methods=['GET'])
@login_required
def get_task_status(task_id):
    """查询指定异步任务的当前状态。"""
    try:
        task_status = task_manager.get_task_status(task_id)
        if not task_status:
            return error_response(404, '任务不存在')
        return success_response(task_status)
    except Exception as e:
        return error_response(500, f'查询任务状态失败: {str(e)}')


@bp.route('/tasks', methods=['GET'])
@login_required
def get_all_tasks():
    """获取所有异步任务列表。"""
    try:
        all_tasks = list(task_manager.tasks.values())
        all_tasks.sort(key=lambda x: x.get('created_at', ''), reverse=True)
        return success_response(all_tasks)
    except Exception as e:
        return error_response(500, f'获取任务列表失败: {str(e)}')


@bp.route('/store-to-kb', methods=['POST'])
@login_required
def store_to_knowledge_base():
    """手动将需求文档内容存入知识库（用户确认后主动触发，避免污染）。"""
    try:
        data = request.get_json() or {}
        content = (data.get('documentContent') or '').strip()
        label = (data.get('label') or '需求文档').strip()

        if not content:
            return error_response(400, '文档内容为空')
        if len(content) < 50:
            return error_response(400, '文档内容过短（至少 50 字），不适合存入知识库')

        _ensure_env_loaded()
        if not os.getenv('EMBEDDING_MODEL'):
            return error_response(400, '知识库未配置（EMBEDDING_MODEL 未设置）')

        from app.services.knowledge_service import upload_document
        filename = f"{label}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        result = upload_document(
            file_content=content.encode('utf-8'),
            filename=filename,
            metadata={'source': 'manual_store', 'label': label, 'user_id': current_user.id},
        )
        return success_response({**result, 'message': f'已存入知识库：{filename}'})
    except (ValueError, RuntimeError) as e:
        return error_response(400, str(e))
    except Exception as e:
        return error_response(500, f'存入知识库失败: {str(e)}')
