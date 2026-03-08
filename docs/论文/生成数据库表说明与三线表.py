# -*- coding: utf-8 -*-
"""
根据毕业论文第四章数据库设计内容，生成 Word 格式的「数据库表说明 + 三线表」。
运行前请安装: pip install python-docx
生成文件: 毕业论文-数据库表说明与三线表.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 表说明与三线表数据（表序号、表中文名、表用途、字段列表：名称, 代码, 主键, 外键, 数据类型）
TABLES = [
    {
        "table_id": "4-1",
        "name": "用户表",
        "name_en": "users",
        "desc": "用户表（users）用于存储系统用户的基本信息，包括登录账号、手机号、邮箱、真实姓名、密码哈希及系统角色等，是身份认证与权限判定的基础表。",
        "rows": [
            ("用户编号", "id", "是", "否", "INT"),
            ("用户名，唯一", "username", "否", "否", "VARCHAR"),
            ("手机号，唯一", "phone", "否", "否", "VARCHAR"),
            ("邮箱", "email", "否", "否", "VARCHAR"),
            ("真实姓名", "real_name", "否", "否", "VARCHAR"),
            ("密码哈希", "password_hash", "否", "否", "VARCHAR"),
            ("系统角色", "role", "否", "否", "ENUM"),
            ("是否激活", "is_active", "否", "否", "BOOLEAN"),
            ("创建时间", "created_at", "否", "否", "TIMESTAMP"),
            ("更新时间", "updated_at", "否", "否", "TIMESTAMP"),
        ],
    },
    {
        "table_id": "4-2",
        "name": "项目表",
        "name_en": "projects",
        "desc": "项目表（projects）用于存储测试项目的基本信息，包括项目名称、描述、状态、负责人及创建者等。",
        "rows": [
            ("项目编号", "id", "是", "否", "INT"),
            ("项目名称，唯一", "project_name", "否", "否", "VARCHAR"),
            ("项目描述", "description", "否", "否", "TEXT"),
            ("项目状态", "status", "否", "否", "ENUM"),
            ("项目负责人ID", "owner_id", "否", "是", "INT"),
            ("创建者ID", "creator_id", "否", "是", "INT"),
            ("是否删除", "is_deleted", "否", "否", "BOOLEAN"),
            ("创建时间", "created_at", "否", "否", "DATETIME"),
            ("更新时间", "updated_at", "否", "否", "DATETIME"),
        ],
    },
    {
        "table_id": "4-3",
        "name": "项目成员表",
        "name_en": "project_members",
        "desc": "项目成员表（project_members）记录项目与用户的多对多关系及成员在项目内的角色。",
        "rows": [
            ("成员ID", "id", "是", "否", "INT"),
            ("项目ID", "project_id", "否", "是", "INT"),
            ("用户ID", "user_id", "否", "是", "INT"),
            ("项目角色", "role", "否", "否", "ENUM"),
            ("加入时间", "joined_at", "否", "否", "DATETIME"),
        ],
    },
    {
        "table_id": "4-4",
        "name": "设备表",
        "name_en": "devices",
        "desc": "设备表（devices）用于存储测试用移动设备的信息，包括设备名称、型号、操作系统类型与版本、设备唯一标识、当前状态及负责人等，供设备管理与任务执行时选用。",
        "rows": [
            ("设备编号", "id", "是", "否", "INT"),
            ("设备名称", "device_name", "否", "否", "VARCHAR"),
            ("设备型号", "device_model", "否", "否", "VARCHAR"),
            ("操作系统类型", "os_type", "否", "否", "ENUM"),
            ("设备唯一标识，唯一", "device_id", "否", "否", "VARCHAR"),
            ("设备状态", "status", "否", "否", "ENUM"),
            ("设备负责人ID", "owner_id", "否", "是", "INT"),
            ("创建时间", "created_at", "否", "否", "TIMESTAMP"),
            ("更新时间", "updated_at", "否", "否", "TIMESTAMP"),
        ],
    },
    {
        "table_id": "4-5",
        "name": "迭代表",
        "name_en": "iterations",
        "desc": "迭代表（iterations）存储项目下的迭代信息及起止日期、状态等。",
        "rows": [
            ("迭代编号", "id", "是", "否", "INT"),
            ("所属项目ID", "project_id", "否", "是", "INT"),
            ("迭代名称", "iteration_name", "否", "否", "VARCHAR"),
            ("迭代状态", "status", "否", "否", "ENUM"),
            ("开始日期", "start_date", "否", "否", "DATETIME"),
            ("结束日期", "end_date", "否", "否", "DATETIME"),
            ("创建时间", "created_at", "否", "否", "TIMESTAMP"),
            ("更新时间", "updated_at", "否", "否", "TIMESTAMP"),
        ],
    },
    {
        "table_id": "4-6",
        "name": "测试套件表",
        "name_en": "test_suites",
        "desc": "测试套件表（test_suites）用于存储用例的树形组织单元，支持文件夹与用例集两种类型，可关联项目、迭代及版本需求，并存储脑图 JSON 数据及版本号。",
        "rows": [
            ("套件编号", "id", "是", "否", "INT"),
            ("套件名称", "suite_name", "否", "否", "VARCHAR"),
            ("父套件ID，树形结构", "parent_id", "否", "是", "INT"),
            ("类型：文件夹/用例集", "type", "否", "否", "ENUM"),
            ("所属项目ID", "project_id", "否", "是", "INT"),
            ("所属迭代ID", "iteration_id", "否", "是", "INT"),
            ("创建者ID", "creator_id", "否", "是", "INT"),
            ("脑图 JSON 数据", "case_mindmap_data", "否", "否", "LONGTEXT"),
            ("脑图版本号", "mindmap_version", "否", "否", "INT"),
            ("创建时间", "created_at", "否", "否", "DATETIME"),
            ("更新时间", "updated_at", "否", "否", "DATETIME"),
        ],
    },
    {
        "table_id": "4-7",
        "name": "测试用例表",
        "name_en": "test_cases",
        "desc": "测试用例表（test_cases）用于存储单条测试用例的标题、优先级、前置条件、测试步骤、预期结果及所属套件、项目等。",
        "rows": [
            ("用例编号", "id", "是", "否", "INT"),
            ("测试用例编号", "case_number", "否", "否", "VARCHAR"),
            ("用例名称", "case_name", "否", "否", "VARCHAR"),
            ("优先级", "priority", "否", "否", "ENUM"),
            ("所属项目ID", "project_id", "否", "是", "INT"),
            ("所属套件ID", "suite_id", "否", "是", "INT"),
            ("前置条件", "preconditions", "否", "否", "TEXT"),
            ("测试步骤", "steps", "否", "否", "TEXT"),
            ("预期结果", "expected_result", "否", "否", "TEXT"),
            ("创建者ID", "creator_id", "否", "是", "INT"),
            ("创建时间", "created_at", "否", "否", "TIMESTAMP"),
            ("更新时间", "updated_at", "否", "否", "TIMESTAMP"),
        ],
    },
    {
        "table_id": "4-8",
        "name": "测试任务表",
        "name_en": "test_tasks",
        "desc": "测试任务表（test_tasks）用于存储测试任务的名称、类型（用例执行/设备脚本）、状态、优先级及关联的套件、创建人、执行人等，是任务执行与报告生成的主表。",
        "rows": [
            ("任务编号", "id", "是", "否", "INT"),
            ("任务名称", "task_name", "否", "否", "VARCHAR"),
            ("任务类型", "task_type", "否", "否", "ENUM"),
            ("任务状态", "status", "否", "否", "ENUM"),
            ("所属项目ID", "project_id", "否", "是", "INT"),
            ("关联测试套件ID", "suite_id", "否", "是", "INT"),
            ("创建者ID", "creator_id", "否", "是", "INT"),
            ("执行者ID", "executor_id", "否", "是", "INT"),
            ("创建时间", "created_at", "否", "否", "TIMESTAMP"),
            ("更新时间", "updated_at", "否", "否", "TIMESTAMP"),
        ],
    },
    {
        "table_id": "4-9",
        "name": "报告表",
        "name_en": "reports",
        "desc": "报告表（reports）用于存储与测试任务关联的报告快照，包括报告类型、任务名称、报告摘要（JSON）、报告明细（JSON）及完成时间等，支持报告详情查询与多格式导出。",
        "rows": [
            ("报告ID", "id", "是", "否", "INT"),
            ("关联任务ID", "task_id", "否", "是", "INT"),
            ("报告类型", "report_type", "否", "否", "ENUM"),
            ("任务名称（冗余）", "task_name", "否", "否", "VARCHAR"),
            ("报告摘要", "summary", "否", "否", "JSON"),
            ("报告明细", "details", "否", "否", "JSON"),
            ("任务完成时间", "completed_at", "否", "否", "TIMESTAMP"),
            ("报告生成时间", "created_at", "否", "否", "TIMESTAMP"),
        ],
    },
    {
        "table_id": "4-10",
        "name": "消息通知表",
        "name_en": "notifications",
        "desc": "消息通知表（notifications）用于存储按用户维度的系统通知，包括类型、标题、摘要、是否已读、关联实体类型与 ID 等，支撑通知列表与实时推送。",
        "rows": [
            ("通知ID", "id", "是", "否", "INT"),
            ("接收人用户ID", "user_id", "否", "是", "INT"),
            ("消息类型", "type", "否", "否", "VARCHAR"),
            ("标题", "title", "否", "否", "VARCHAR"),
            ("摘要", "summary", "否", "否", "TEXT"),
            ("是否已读", "is_read", "否", "否", "BOOLEAN"),
            ("关联实体类型", "related_type", "否", "否", "VARCHAR"),
            ("关联实体ID", "related_id", "否", "否", "INT"),
            ("创建时间", "created_at", "否", "否", "DATETIME"),
        ],
    },
]


def set_cell_border(cell, **kwargs):
    """设置单元格边框。未传入的边设为 nil（无边框）。kwargs 键: start, top, end, bottom, insideH, insideV；值: {"val": "single", "sz": 12}。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 移除已有 tcBorders
    for b in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(b)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if edge in kwargs:
            el.set(qn("w:val"), kwargs[edge].get("val", "single"))
            el.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
        else:
            el.set(qn("w:val"), "nil")
        el.set(qn("w:space"), "0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def apply_three_line_table(table):
    """为表格应用三线表样式：顶线粗、表头下划线、底线粗，无竖线、无内部横线。"""
    sz_thick = 12   # 1.5 磅
    sz_thin = 4     # 0.5 磅
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0:
                set_cell_border(
                    cell,
                    top={"val": "single", "sz": sz_thick},
                    bottom={"val": "single", "sz": sz_thin},
                )
            elif row_idx == len(table.rows) - 1:
                set_cell_border(
                    cell,
                    bottom={"val": "single", "sz": sz_thick},
                )
            else:
                set_cell_border(cell)


def main():
    doc = Document()
    # 正文宋体、小四
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    for idx, t in enumerate(TABLES):
        table_id = t["table_id"]
        name = t["name"]
        desc = t["desc"]
        rows = t["rows"]

        # 表前说明：(序号)表名: 用途,如下表X所示。
        run_desc = f"（{idx + 1}）{name}：{desc.strip().rstrip('。')}，如下表{table_id}所示。"
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12) if idx > 0 else Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(run_desc)

        # 表标题：表X 表名
        p_cap = doc.add_paragraph()
        p_cap.add_run(f"表 {table_id}  {name}").bold = True
        p_cap.paragraph_format.space_after = Pt(6)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 三线表：名称、代码、主键、外键、数据类型
        col_count = 5
        table = doc.add_table(rows=1 + len(rows), cols=col_count)
        table.style = "Table Grid"  # 基础样式，三线由下方 apply_three_line_table 覆盖边框
        header_cells = table.rows[0].cells
        headers = ["名称", "代码", "主键", "外键", "数据类型"]
        for i, h in enumerate(headers):
            header_cells[i].text = h
            for p in header_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10.5)
                    r.font.name = "宋体"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        for r_idx, row_data in enumerate(rows):
            row_cells = table.rows[r_idx + 1].cells
            for c_idx, val in enumerate(row_data):
                row_cells[c_idx].text = str(val)
                for p in row_cells[c_idx].paragraphs:
                    if c_idx in (2, 3):  # 主键、外键居中
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.size = Pt(10.5)
                        r.font.name = "宋体"
                        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        apply_three_line_table(table)
        doc.add_paragraph()  # 表后空一行

    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "毕业论文-数据库表说明与三线表.docx")
    doc.save(out_path)
    print("已生成: " + out_path)
    return out_path


if __name__ == "__main__":
    main()
