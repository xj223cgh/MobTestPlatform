# -*- coding: utf-8 -*-
"""
将数据库总览表（表4-0）写入 Word 文档「毕业论文-数据库表说明与三线表.docx」。
表头：表名、中文名称、说明。表格采用三线表样式（顶线、表头下划线、底线）。
运行前请安装：pip install python-docx
"""
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 文档路径（与本脚本同目录）
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(BASE_DIR, '毕业论文-数据库表说明与三线表.docx')

# 数据库总览表数据：按功能模块、优先级排序 (表名, 中文名称, 说明)
# 顺序：用户与权限 → 系统配置 → 项目管理 → 设备管理 → 测试套件与用例 → 任务与执行 → 报告 → 通知
OVERVIEW_ROWS = [
    # 用户与权限
    ('users', '用户表', '存储系统用户基本信息，是身份认证与权限判定的基础表'),
    ('email_verify_codes', '邮箱验证码表', '邮箱验证码（登录验证码，5分钟有效），支撑登录验证'),
    ('role_permissions', '角色权限表', '角色与功能埋点配置，支撑权限判定'),
    # 系统配置
    ('system_settings', '系统设置表', '系统全局设置，键值对存储'),
    ('user_settings', '用户设置表', '用户个人设置'),
    # 项目管理（含迭代与需求）
    ('projects', '项目表', '存储测试项目基本信息'),
    ('project_members', '项目成员表', '项目与用户的多对多关系及成员在项目内的角色'),
    ('iterations', '迭代表', '迭代信息，关联项目'),
    ('version_requirements', '版本需求表', '版本需求信息，关联项目与迭代'),
    # 设备管理
    ('agents', 'Agent表', '本机Agent（每台管理USB设备的电脑运行的Agent程序）'),
    ('user_agent_bindings', '用户Agent绑定表', '用户与Agent的绑定关系，一个用户当前只绑定一台本机Agent'),
    ('agent_binding_codes', 'Agent绑定码表', '绑定码，供一键绑定或Agent输入完成绑定，短期有效'),
    ('devices', '设备表', '存储测试用移动设备信息'),
    # 测试套件与用例
    ('test_suites', '测试套件表', '测试套件/用例集，含脑图数据与评审状态'),
    ('mindmap_versions', '脑图版本表', '脑图编辑版本快照，用于回退'),
    ('test_suite_review_tasks', '用例集评审任务表', '用例集评审任务'),
    ('test_case_review_details', '用例评审详情表', '单条用例评审详情，关联评审任务与用例'),
    ('test_suite_review_history', '用例集评审历史表', '评审历史记录'),
    ('test_case_review_history', '用例评审历史表', '用例评审历史记录'),
    ('test_cases', '测试用例表', '测试用例'),
    ('case_tags', '用例标签表', '用例标签字典，按项目维护'),
    ('case_markers', '用例标记表', '用例标记字典，按项目维护'),
    # 任务与执行
    ('test_tasks', '测试任务表', '测试任务'),
    ('task_folders', '任务文件夹表', '任务文件夹，按任务类型分组'),
    ('task_case_relation', '任务-用例关联表', '测试任务与测试用例的多对多关联'),
    ('task_device_relation', '任务-设备关联表', '测试任务与设备的多对多关联'),
    ('task_case_snapshots', '任务用例快照表', '任务关联用例快照，用于历史追溯'),
    ('test_case_executions', '用例执行结果表', '测试用例执行结果，关联任务与迭代'),
    # 报告
    ('reports', '报告表', '任务执行完成时的数据快照，列表与详情以快照为准'),
    # 通知
    ('notifications', '通知表', '消息通知，按接收人维度存储'),
]


def set_cell_border(cell, **kwargs):
    """设置单元格边框。kwargs: top, bottom, start, end 等，值为 True 表示保留，False 表示去掉。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if kwargs.get(edge, True):
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def set_table_three_line_style(table, header_row_index=0):
    """三线表：仅保留顶线、表头下边线、底线。"""
    def get_tcPr(cell):
        return cell._tc.get_or_add_tcPr()

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = get_tcPr(cell)
            # 移除已有边框
            for prev in tcPr.findall('.//' + qn('w:tcBorders')):
                tcPr.remove(prev)
            tcBorders = OxmlElement('w:tcBorders')
            # 顶线：仅第一行
            if row_idx == 0:
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '12')  # 稍粗
                top.set(qn('w:color'), '000000')
                tcBorders.append(top)
            # 底线：表头行下方 或 最后一行
            if row_idx == header_row_index or row_idx == len(table.rows) - 1:
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12' if (row_idx == header_row_index or row_idx == len(table.rows) - 1) else '4')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
            # 左、右边线不显示（三线表通常无竖线）
            if len(tcBorders) > 0:
                tcPr.append(tcBorders)


def main():
    if not os.path.exists(DOCX_PATH):
        doc = Document()
        doc.add_paragraph('数据库表说明与三线表')
        doc.add_paragraph()
    else:
        doc = Document(DOCX_PATH)

    # 在文档最前面插入：标题 + 总览表
    body = doc.element.body
    # 标题段落
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('表4-0  数据库总览表')
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = 1  # 1 = CENTER
    # 表格：1 表头行 + 30 数据行
    table = doc.add_table(rows=1 + len(OVERVIEW_ROWS), cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = '表名'
    header_cells[1].text = '中文名称'
    header_cells[2].text = '说明'
    for c in header_cells:
        for para in c.paragraphs:
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
                if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for i, (name, cn_name, desc) in enumerate(OVERVIEW_ROWS):
        row = table.rows[i + 1]
        row.cells[0].text = name
        row.cells[1].text = cn_name
        row.cells[2].text = desc
        for c in row.cells:
            for para in c.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_table_three_line_style(table, header_row_index=0)
    # 在正文最前插入：先插表格，再插段落，这样段落在表格上方
    body.insert(0, table._tbl)
    body.insert(0, p._element)
    # 在总览表后加一空行
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    body.insert(2, p2._element)

    doc.save(DOCX_PATH)
    print('已写入：', DOCX_PATH)
    print('表4-0  数据库总览表（共 %d 行数据）已插入到文档开头。' % len(OVERVIEW_ROWS))


if __name__ == '__main__':
    main()
